import os
import sys
import re
import json
import copy
import random
import threading
from datetime import datetime, timedelta
import customtkinter as ctk
from tkinter import filedialog, messagebox
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _resource_path(*parts):
    """Caminho para arquivo somente-leitura (bundled no _MEIPASS quando frozen)."""
    if getattr(sys, 'frozen', False):
        base = sys._MEIPASS
    else:
        base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(base, *parts)


def _writable_path(*parts):
    """Caminho para arquivo gravável (pasta ao lado do .exe quando frozen)."""
    if getattr(sys, 'frozen', False):
        base = os.path.dirname(sys.executable)
    else:
        base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(base, *parts)


_TEMPLATE_PROFESSOR          = _resource_path("arquivos", "Relatorio Professor.docx")
_TEMPLATE_OUTRO_CARGOS       = _resource_path("arquivos", "Relatorio Outro Cargos.docx")
_TEMPLATE_CTS                = _resource_path("arquivos", "Relatorio ctc.doc")
_TEMPLATE_ESTAGIO            = _resource_path("arquivos", "Relatorio estagio.doc")
_HISTORICO_PATH              = _writable_path("arquivos", "sugestoes_professor.json")
_HISTORICO_OUTRO_CARGOS_PATH = _writable_path("arquivos", "sugestoes_outro_cargos.json")
_CARGOS_PATH                 = _resource_path("arquivos", "cargos.json")
_ASSINATURAS_PATH            = _resource_path("arquivos", "assinatura.json")
_ATRIBUICOES_PATH             = _resource_path("arquivos", "atribuicoes.json")


# ---------------------------------------------------------------------------
# Histórico de sugestões (CURSO / INSTITUICAO)
# ---------------------------------------------------------------------------
def _carregar_historico(path=None) -> dict:
    if path is None:
        path = _HISTORICO_PATH
    if os.path.exists(path):
        try:
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return {}


def _salvar_historico(data: dict, path=None):
    if path is None:
        path = _HISTORICO_PATH
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def _carregar_cargos() -> dict:
    if os.path.exists(_CARGOS_PATH):
        try:
            with open(_CARGOS_PATH, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return {}


def _carregar_assinaturas() -> list[dict]:
    if os.path.exists(_ASSINATURAS_PATH):
        try:
            with open(_ASSINATURAS_PATH, "r", encoding="utf-8") as f:
                dados = json.load(f)
            return [item for item in dados if item.get("nome") and item.get("portaria")]
        except (OSError, json.JSONDecodeError, AttributeError):
            pass
    return []


def _carregar_atribuicoes() -> list[dict]:
    if os.path.exists(_ATRIBUICOES_PATH):
        try:
            with open(_ATRIBUICOES_PATH, "r", encoding="utf-8") as f:
                dados = json.load(f)
            return sorted(
                [item for item in dados if item.get("nome") and item.get("descricao")],
                key=lambda item: item["nome"].casefold(),
            )
        except (OSError, json.JSONDecodeError, AttributeError):
            pass
    return []


def _bind_scroll_dropdown(widget: ctk.CTkOptionMenu, valores: list):
    # O DropdownMenu do CTkOptionMenu é um tkinter.Menu nativo do Windows.
    # Quando aberto via post(), o OS intercepta todos os eventos de mouse
    # (TrackPopupMenu tem seu próprio loop de mensagens) — impossível capturar
    # <MouseWheel> via Tkinter enquanto o menu nativo estiver visível.
    # O scroll funciona apenas com o dropdown fechado.
    def _on_scroll(event):
        try:
            idx = valores.index(widget.get())
        except ValueError:
            return "break"
        direction = 1 if event.delta < 0 else -1
        novo_idx = max(0, min(len(valores) - 1, idx + direction))
        if novo_idx != idx:
            novo_val = valores[novo_idx]
            widget.set(novo_val)
            cmd = widget._command
            if cmd:
                cmd(novo_val)
        return "break"

    _sf = None
    _p = getattr(widget, 'master', None)
    for _ in range(10):
        if _p is None:
            break
        if isinstance(_p, ctk.CTkScrollableFrame):
            _sf = _p
            break
        _p = getattr(_p, 'master', None)

    if _sf is None:
        return

    _canvas = _sf._parent_canvas  # tkinter.Canvas — bind_all sem restrição CTk

    def _still_over_widget(event):
        try:
            return (widget.winfo_rootx() <= event.x_root <= widget.winfo_rootx() + widget.winfo_width()
                    and widget.winfo_rooty() <= event.y_root <= widget.winfo_rooty() + widget.winfo_height())
        except Exception:
            return True

    def _on_enter(_):
        _canvas.bind_all("<MouseWheel>", _on_scroll)

    def _on_leave(event):
        if _still_over_widget(event):
            return
        if hasattr(_sf, '_mouse_wheel_all'):
            _canvas.bind_all("<MouseWheel>", _sf._mouse_wheel_all)

    widget.bind("<Enter>", _on_enter, add="+")
    widget.bind("<Leave>", _on_leave, add="+")


# ---------------------------------------------------------------------------
# Mapeamentos e definição dos campos — Professor
# ---------------------------------------------------------------------------
_GENERO_MAP = {
    "Masculino": "o servidor",
    "Feminino":  "a servidora",
}

_PORCENTAGEM_MAP  = {"1": "15%",           "2": "25%",      "3": "40%"}
_NIVEL_CURSO_MAP  = {"1": "Especialização", "2": "Mestrado",  "3": "Doutorado"}

_MESES_PT = [
    "Janeiro", "Fevereiro", "Março",    "Abril",   "Maio",      "Junho",
    "Julho",   "Agosto",    "Setembro", "Outubro", "Novembro",  "Dezembro",
]
_DIAS = [str(d) for d in range(1, 32)]
_ANOS = [str(a) for a in range(1990, 2031)]

_hoje     = datetime.now()
_DIA_ATUAL = str(_hoje.day)
_MES_ATUAL = _MESES_PT[_hoje.month - 1]
_ANO_ATUAL = str(_hoje.year)


def _data_por_extenso(data):
    return f"{data.day} de {_MESES_PT[data.month - 1].lower()} de {data.year}"


def _numero_por_extenso(numero):
    unidades = ["zero", "um", "dois", "três", "quatro", "cinco", "seis", "sete", "oito", "nove"]
    especiais = {10: "dez", 11: "onze", 12: "doze", 13: "treze", 14: "quatorze", 15: "quinze", 16: "dezesseis", 17: "dezessete", 18: "dezoito", 19: "dezenove"}
    dezenas = ["", "", "vinte", "trinta", "quarenta", "cinquenta", "sessenta", "setenta", "oitenta", "noventa"]
    centenas = ["", "cento", "duzentos", "trezentos", "quatrocentos", "quinhentos", "seiscentos", "setecentos", "oitocentos", "novecentos"]
    if numero < 10:
        return unidades[numero]
    if numero < 20:
        return especiais[numero]
    if numero < 100:
        return dezenas[numero // 10] + (f" e {unidades[numero % 10]}" if numero % 10 else "")
    if numero == 100:
        return "cem"
    if numero < 1000:
        return centenas[numero // 100] + (f" e {_numero_por_extenso(numero % 100)}" if numero % 100 else "")
    if numero < 1000000:
        milhares, resto = divmod(numero, 1000)
        prefixo = "mil" if milhares == 1 else f"{_numero_por_extenso(milhares)} mil"
        return prefixo + (f" {_numero_por_extenso(resto)}" if resto else "")
    milhoes, resto = divmod(numero, 1000000)
    prefixo = "um milhão" if milhoes == 1 else f"{_numero_por_extenso(milhoes)} milhões"
    return prefixo + (f" {_numero_por_extenso(resto)}" if resto else "")

# Formato: (tipo, placeholder, label, extra)
#   tipo "entry"        → campo de texto livre
#   tipo "dropdown"     → CTkOptionMenu; extra = lista de opções
#   tipo "dropdown_map" → CTkOptionMenu com mapeamento; extra = dict {exibição: valor}
#   tipo "auto"         → calculado automaticamente, sem input do usuário
_CAMPOS_PROFESSOR = [
    ("Documento", [
        ("entry",        "{{NUMERO_DOC}}", "Número do Documento",     None),
    ]),
    ("Dados do Servidor", [
        ("entry",        "{{NOME}}",        "Nome",                    None),
        ("dropdown_map", "{{GENERO}}",      "Gênero",                  _GENERO_MAP),
        ("dropdown",     "{{NIVEL}}",       "Nível Atual",             ["1", "2", "3"]),
        ("date",         "{{DATA_INICIO}}", "Data de Início no Cargo", None),
    ]),
    ("Alteração de Nível", [
        ("auto",         "{{NIVEL+1}}",     "Próximo Nível",           None),
        ("auto",         "{{PORCENTAGEM}}", "Porcentagem do Acréscimo",None),
        ("auto",         "{{NIVEL_CURSO}}", "Tipo de Formação",        None),
        ("combobox",     "{{CURSO}}",       "Nome do Curso",           None),
        ("combobox",     "{{INSTITUICAO}}", "Instituição",             None),
    ]),
    ("Data do Parecer", [
        ("dropdown", "{{DATA_DIA}}", "Dia", (_DIAS,     _DIA_ATUAL)),
        ("dropdown", "{{DATA_MES}}", "Mês", (_MESES_PT, _MES_ATUAL)),
        ("dropdown", "{{DATA_ANO}}", "Ano", (_ANOS,     _ANO_ATUAL)),
    ]),
]

# Campos obrigatórios por tipo
_ENTRY_PLACEHOLDERS = {
    ph for _, campos in _CAMPOS_PROFESSOR
    for tipo, ph, *_ in campos
    if tipo in ("entry", "combobox")
}
_DATE_PLACEHOLDERS = {
    ph for _, campos in _CAMPOS_PROFESSOR
    for tipo, ph, *_ in campos
    if tipo == "date"
}


# ---------------------------------------------------------------------------
# Widget: entrada com máscara DD/MM/AAAA
# ---------------------------------------------------------------------------
class _DateEntry(ctk.CTkEntry):
    """CTkEntry com máscara automática DD/MM/AAAA e validação de data real."""

    def __init__(self, master, **kwargs):
        self._var = ctk.StringVar()
        super().__init__(master, textvariable=self._var, placeholder_text="DD/MM/AAAA", **kwargs)
        self._updating = False
        self._var.trace_add("write", self._formatar)

    def _formatar(self, *_):
        if self._updating:
            return
        self._updating = True
        raw = self._var.get()
        digits = re.sub(r"\D", "", raw)[:8]
        formatted = digits[:2]
        if len(digits) >= 3:
            formatted += "/" + digits[2:4]
        if len(digits) >= 5:
            formatted += "/" + digits[4:8]
        if formatted != raw:
            self._var.set(formatted)
            self.after(0, lambda: self.icursor("end"))
        self._updating = False

    def get(self) -> str:
        return self._var.get()

    def is_valid(self) -> bool:
        try:
            datetime.strptime(self._var.get(), "%d/%m/%Y")
            return True
        except ValueError:
            return False


# ---------------------------------------------------------------------------
# Widget: entrada numérica de até 4 dígitos (matrícula)
# ---------------------------------------------------------------------------
class _MatriculaEntry(ctk.CTkEntry):

    def __init__(self, master, **kwargs):
        self._var = ctk.StringVar()
        super().__init__(master, textvariable=self._var, placeholder_text="0000", **kwargs)
        self._updating = False
        self._var.trace_add("write", self._filtrar)

    def _filtrar(self, *_):
        if self._updating:
            return
        self._updating = True
        raw = self._var.get()
        digits = re.sub(r"\D", "", raw)[:4]
        if digits != raw:
            self._var.set(digits)
            self.after(0, lambda: self.icursor("end"))
        self._updating = False

    def get(self) -> str:
        return self._var.get()

    def is_valid(self) -> bool:
        v = self._var.get()
        return bool(v) and v.isdigit()


class _LettersEntry(ctk.CTkEntry):

    def __init__(self, master, **kwargs):
        self._var = ctk.StringVar()
        super().__init__(master, textvariable=self._var, **kwargs)
        self._updating = False
        self._var.trace_add("write", self._filtrar)

    def _filtrar(self, *_):
        if self._updating:
            return
        self._updating = True
        valor = re.sub(r"[^\w\sÀ-ÿ]", "", self._var.get(), flags=re.UNICODE).replace("_", "")
        if valor != self._var.get():
            self._var.set(valor)
            self.after(0, lambda: self.icursor("end"))
        self._updating = False

    def get(self) -> str:
        return self._var.get()


class _DigitsEntry(ctk.CTkEntry):

    def __init__(self, master, max_length=None, **kwargs):
        self._var = ctk.StringVar()
        self._max_length = max_length
        super().__init__(master, textvariable=self._var, **kwargs)
        self._updating = False
        self._var.trace_add("write", self._filtrar)

    def _filtrar(self, *_):
        if self._updating:
            return
        self._updating = True
        valor = re.sub(r"\D", "", self._var.get())
        if self._max_length:
            valor = valor[:self._max_length]
        if valor != self._var.get():
            self._var.set(valor)
            self.after(0, lambda: self.icursor("end"))
        self._updating = False

    def get(self) -> str:
        return self._var.get()


class _CpfEntry(_DigitsEntry):

    def _filtrar(self, *_):
        if self._updating:
            return
        self._updating = True
        digits = re.sub(r"\D", "", self._var.get())[:11]
        valor = digits
        if len(digits) > 3:
            valor = digits[:3] + "." + digits[3:]
        if len(digits) > 6:
            valor = valor[:7] + "." + valor[7:]
        if len(digits) > 9:
            valor = valor[:11] + "-" + valor[11:]
        if valor != self._var.get():
            self._var.set(valor)
            self.after(0, lambda: self.icursor("end"))
        self._updating = False

    def is_valid(self):
        return len(re.sub(r"\D", "", self._var.get())) == 11


def _preencher_entry(widget, valor):
    if isinstance(widget, ctk.CTkComboBox):
        widget.set(valor)
    else:
        widget.delete(0, "end")
        widget.insert(0, valor)


def _preencher_widgets_teste(widget, indice=0):
    for filho in widget.winfo_children():
        if isinstance(filho, (ctk.CTkEntry, ctk.CTkComboBox)):
            _preencher_entry(filho, f"Teste {indice + 1}")
            indice += 1
        elif isinstance(filho, ctk.CTkOptionMenu):
            valores = filho.cget("values")
            if valores:
                filho.set(random.choice(valores))
        indice = _preencher_widgets_teste(filho, indice)
    return indice


def _listar_entries(widget):
    encontrados = []
    for filho in widget.winfo_children():
        if isinstance(filho, (ctk.CTkEntry, ctk.CTkComboBox)):
            encontrados.append(filho)
        encontrados.extend(_listar_entries(filho))
    return encontrados


def _definir_data_teste(widgets, data):
    widgets[0].set(str(data.day))
    widgets[1].set(_MESES_PT[data.month - 1])
    widgets[2].set(str(data.year))


def _dados_pessoais_teste():
    nomes = ("Rodrigo Santos", "Mariana Oliveira", "Carlos Pereira", "Ana Martins")
    return random.choice(nomes), str(random.randint(10000000, 99999999)), str(random.randint(10000000000, 99999999999))


def _local_teste(indice):
    return f"Escola {random.choice(('Modelo', 'Municipal', 'Central', 'Horizonte'))} {indice + 1}"


# ---------------------------------------------------------------------------
# Página principal de Documentação
# ---------------------------------------------------------------------------
class DocumentacaoPage(ctk.CTkFrame):
    page_key = "documentacao"

    _TIPOS = ["Professor", "Outros Cargos", "Certidões (CTS)", "Estagiário (CTS)"]

    def __init__(self, parent):
        super().__init__(parent, corner_radius=0)
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)
        self._build_ui()

    def _build_ui(self):
        header = ctk.CTkFrame(self, height=118, corner_radius=0, fg_color=("gray88", "gray15"))
        header.grid(row=0, column=0, sticky="ew")
        header.grid_columnconfigure(0, weight=1)
        header.grid_rowconfigure(1, weight=1)
        self._header = header

        ctk.CTkLabel(
            header, text="Documentação", font=("Arial", 20, "bold")
        ).grid(row=0, column=0, padx=25, pady=(10, 4), sticky="w")

        controles = ctk.CTkFrame(header, fg_color="transparent")
        controles.grid(row=1, column=0, padx=25, pady=(0, 8), sticky="ew")
        self._controles_header = controles

        assinaturas = _carregar_assinaturas()
        self._assinaturas_por_nome = {item["nome"]: item for item in assinaturas}
        nomes_assinaturas = list(self._assinaturas_por_nome) or ["— Nenhuma assinatura cadastrada —"]
        self._assinatura_nome = ctk.CTkOptionMenu(controles, values=nomes_assinaturas, width=190)
        self._assinatura_cargo = ctk.CTkOptionMenu(
            controles,
            values=[
                "Diretor de Recursos Humanos",
                "Oficial Administrativo",
                "Chefe de Divisão",
                "Assessor de Recursos Humanos",
                "Auxiliar Administrativo"
            ],
            width=190,
        )

        self._tipo_label = ctk.CTkLabel(
            controles, text="Tipo de documento:", font=("Arial", 12), text_color="gray55"
        )

        self._dropdown = ctk.CTkOptionMenu(
            controles,
            values=self._TIPOS,
            width=180,
            command=self._on_tipo_alterado,
        )
        self._tipo_atual = self._TIPOS[0]
        self._teste_button = ctk.CTkButton(
            controles, text="Teste", width=90, command=self._preencher_teste
        )
        self._header_widgets = [
            self._assinatura_nome,
            self._assinatura_cargo,
            self._tipo_label,
            self._dropdown,
            self._teste_button,
        ]
        self._reorganizar_cabecalho()
        header.bind("<Configure>", lambda _event: self._reorganizar_cabecalho(), add="+")

        self._content_area = ctk.CTkFrame(self, fg_color="transparent")
        self._content_area.grid(row=1, column=0, sticky="nsew")
        self._content_area.grid_columnconfigure(0, weight=1)
        self._content_area.grid_rowconfigure(0, weight=1)

        self._sub_paginas = {
            "Professor":     _SubPaginaProfessor(self._content_area),
            "Outros Cargos": _SubPaginaOutrosCargos(self._content_area),
            "Certidões (CTS)": _SubPaginaCTS(self._content_area),
            "Estagiário (CTS)": _SubPaginaEstagiario(self._content_area),
        }
        for sub in self._sub_paginas.values():
            sub._obter_assinatura = self._obter_assinatura
            sub.grid(row=0, column=0, sticky="nsew")

        self._mostrar("Professor")

    def _reorganizar_cabecalho(self):
        largura = self._header.winfo_width()
        if largura <= 1:
            return
        for widget in self._header_widgets:
            widget.grid_forget()
        for coluna in range(5):
            self._controles_header.grid_columnconfigure(coluna, weight=1)

        if largura < 620:
            for row, widget in enumerate(self._header_widgets):
                widget.grid(row=row, column=0, padx=4, pady=2, sticky="ew")
            altura = 260
        elif largura < 980:
            for coluna, widget in enumerate(self._header_widgets[:2]):
                widget.grid(row=0, column=coluna, padx=4, pady=2, sticky="ew")
            self._tipo_label.grid(row=1, column=0, padx=4, pady=2, sticky="e")
            self._dropdown.grid(row=1, column=1, padx=4, pady=2, sticky="ew")
            self._teste_button.grid(row=2, column=0, columnspan=2, padx=4, pady=2, sticky="ew")
            altura = 150
        else:
            for coluna, widget in enumerate(self._header_widgets):
                widget.grid(row=0, column=coluna, padx=4, pady=2, sticky="ew")
            altura = 78
        self._header.configure(height=altura)

    def _on_tipo_alterado(self, valor):
        self._tipo_atual = valor
        self._mostrar(valor)

    def _preencher_teste(self):
        subpagina = self._sub_paginas[self._tipo_atual]
        preencher = getattr(subpagina, "preencher_teste", None)
        if preencher:
            preencher()

    def _obter_assinatura(self):
        nome = self._assinatura_nome.get()
        dados = self._assinaturas_por_nome.get(nome, {})
        return {
            "{{NOME_ASSINATURA}}": dados.get("nome", ""),
            "{{PORTARIA ASSINATURA}}": dados.get("portaria", ""),
            "{{CARGO ASSINATURA}}": self._assinatura_cargo.get(),
        }

    def _mostrar(self, tipo):
        for sub in self._sub_paginas.values():
            sub.grid_remove()
        self._sub_paginas[tipo].grid()


# ---------------------------------------------------------------------------
# Sub-página: Professor
# ---------------------------------------------------------------------------
class _SubPaginaProfessor(ctk.CTkFrame):

    def __init__(self, parent):
        super().__init__(parent, corner_radius=0, fg_color="transparent")
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)

        self._getters: dict[str, callable] = {}
        self._date_widgets: dict[str, _DateEntry] = {}
        self._combos: dict[str, ctk.CTkComboBox] = {}
        self._nivel_auto_label: ctk.CTkLabel | None = None
        self._porcentagem_auto_label: ctk.CTkLabel | None = None
        self._nivel_curso_auto_label: ctk.CTkLabel | None = None

        self._build_toolbar()
        self._build_form()
        self._build_statusbar()

    # ------------------------------------------------------------------ UI

    def _build_toolbar(self):
        bar = ctk.CTkFrame(self, height=55, fg_color=("gray82", "gray18"))
        bar.grid(row=0, column=0, sticky="ew")
        bar.grid_propagate(False)

        self._progress = ctk.CTkProgressBar(bar, mode="indeterminate", width=120, height=8)

        ctk.CTkButton(
            bar, text="💾 Exportar .docx", width=150, command=self._exportar
        ).pack(side="right", padx=15, pady=10)

    def _build_form(self):
        scroll = ctk.CTkScrollableFrame(self, fg_color="transparent")
        scroll.grid(row=1, column=0, sticky="nsew", padx=24, pady=(16, 0))
        scroll.grid_columnconfigure(1, weight=1)

        row = 0
        for grupo, campos in _CAMPOS_PROFESSOR:
            ctk.CTkLabel(
                scroll,
                text=grupo.upper(),
                font=("Arial", 10, "bold"),
                text_color="gray50",
            ).grid(row=row, column=0, columnspan=2, sticky="w", pady=(16, 4))
            row += 1

            for tipo, ph, label, extra in campos:
                ctk.CTkLabel(
                    scroll, text=label, font=("Arial", 12), anchor="w"
                ).grid(row=row, column=0, sticky="w", padx=(8, 16), pady=4)

                if tipo == "entry":
                    widget = ctk.CTkEntry(scroll, height=32, font=("Arial", 12))
                    widget.grid(row=row, column=1, sticky="ew", pady=4)
                    self._getters[ph] = widget.get

                elif tipo == "date":
                    widget = _DateEntry(scroll, height=32, font=("Arial", 12))
                    widget.grid(row=row, column=1, sticky="ew", pady=4)
                    self._getters[ph] = widget.get
                    self._date_widgets[ph] = widget

                elif tipo == "combobox":
                    historico = _carregar_historico()
                    sugestoes = historico.get(ph, [])
                    widget = ctk.CTkComboBox(
                        scroll, values=sugestoes, height=32, font=("Arial", 12)
                    )
                    widget.set("")
                    widget.grid(row=row, column=1, sticky="ew", pady=4)
                    self._getters[ph] = widget.get
                    self._combos[ph] = widget

                elif tipo == "dropdown":
                    opcoes, default = extra if isinstance(extra, tuple) else (extra, extra[0])
                    widget = ctk.CTkOptionMenu(
                        scroll, values=opcoes, height=32, font=("Arial", 12)
                    )
                    widget.set(default)
                    widget.grid(row=row, column=1, sticky="ew", pady=4)
                    self._getters[ph] = widget.get
                    if ph == "{{NIVEL}}":
                        widget.configure(command=self._on_nivel_alterado)
                    _bind_scroll_dropdown(widget, opcoes)

                elif tipo == "dropdown_map":
                    opcoes = list(extra.keys())
                    widget = ctk.CTkOptionMenu(
                        scroll, values=opcoes, height=32, font=("Arial", 12)
                    )
                    widget.set(opcoes[0])
                    widget.grid(row=row, column=1, sticky="ew", pady=4)
                    mapa = extra
                    self._getters[ph] = lambda w=widget, m=mapa: m[w.get()]
                    _bind_scroll_dropdown(widget, opcoes)

                elif tipo == "auto":
                    _iniciais = {
                        "{{NIVEL+1}}":     "2",
                        "{{PORCENTAGEM}}": _PORCENTAGEM_MAP["1"],
                        "{{NIVEL_CURSO}}": _NIVEL_CURSO_MAP["1"],
                    }
                    lbl = ctk.CTkLabel(
                        scroll,
                        text=_iniciais.get(ph, ""),
                        font=("Arial", 12),
                        text_color="gray55",
                        anchor="w",
                    )
                    lbl.grid(row=row, column=1, sticky="w", padx=6, pady=4)

                    if ph == "{{NIVEL+1}}":
                        self._nivel_auto_label = lbl
                    elif ph == "{{PORCENTAGEM}}":
                        self._porcentagem_auto_label = lbl
                    elif ph == "{{NIVEL_CURSO}}":
                        self._nivel_curso_auto_label = lbl

                    self._getters[ph] = lambda w=lbl: w.cget("text")

                row += 1

    def _build_statusbar(self):
        self._status = ctk.CTkLabel(
            self, text="", font=("Arial", 11), text_color="gray55"
        )
        self._status.grid(row=2, column=0, padx=25, pady=(4, 8), sticky="w")

    # ------------------------------------------------------------------ Callbacks

    def _on_nivel_alterado(self, valor: str):
        if self._nivel_auto_label:
            self._nivel_auto_label.configure(text=str(int(valor) + 1))
        if self._porcentagem_auto_label:
            self._porcentagem_auto_label.configure(text=_PORCENTAGEM_MAP[valor])
        if self._nivel_curso_auto_label:
            self._nivel_curso_auto_label.configure(text=_NIVEL_CURSO_MAP[valor])

    def preencher_teste(self):
        _preencher_widgets_teste(self)
        entries = _listar_entries(self)
        nome, _, _ = _dados_pessoais_teste()
        valores = (str(random.randint(1000, 9999)), nome, "Administração", _local_teste(0))
        for indice, valor in enumerate(valores):
            if indice < len(entries):
                _preencher_entry(entries[indice], valor)

    # ------------------------------------------------------------------ Export

    def _exportar(self):
        if not os.path.exists(_TEMPLATE_PROFESSOR):
            messagebox.showerror("Erro", f"Template não encontrado:\n{_TEMPLATE_PROFESSOR}")
            return

        substituicoes = {ph: getter() for ph, getter in self._getters.items()}
        substituicoes.update(getattr(self, "_obter_assinatura", lambda: {})())

        vazios = [ph for ph in _ENTRY_PLACEHOLDERS if not substituicoes.get(ph, "").strip()]
        datas_invalidas = [ph for ph, w in self._date_widgets.items() if not w.is_valid()]
        problemas = vazios + datas_invalidas
        if problemas:
            labels_problema = [
                label + (" (data inválida)" if ph in datas_invalidas else "")
                for _, campos in _CAMPOS_PROFESSOR
                for tipo, ph, label, *_ in campos
                if ph in problemas
            ]
            messagebox.showwarning(
                "Campos inválidos",
                "Corrija os campos antes de exportar:\n• " + "\n• ".join(labels_problema),
            )
            return

        destino = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Documento Word", "*.docx")],
            initialfile="Relatorio Professor_editado.docx",
        )
        if not destino:
            return

        self._set_loading(True)
        threading.Thread(
            target=self._thread_exportar,
            args=(_TEMPLATE_PROFESSOR, substituicoes, destino),
            daemon=True,
        ).start()

    def _thread_exportar(self, origem, subs, destino):
        try:
            doc = Document(origem)
            for para in doc.paragraphs:
                _substituir_paragrafo(para, subs)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            _substituir_paragrafo(para, subs)
            doc.save(destino)
            self.after(0, lambda: self._on_exportado(destino))
        except Exception as e:
            msg = str(e)
            self.after(0, lambda: self._on_erro(msg))

    def _on_exportado(self, destino):
        self._set_loading(False)
        self._set_status(f"Exportado com sucesso: {os.path.basename(destino)}", "success")
        self.after(200, self._verificar_novas_sugestoes)

    def _verificar_novas_sugestoes(self):
        historico = _carregar_historico()
        novos = {}
        for ph, combo in self._combos.items():
            valor = combo.get().strip()
            if valor and valor not in historico.get(ph, []):
                label = next(
                    label for _, campos in _CAMPOS_PROFESSOR
                    for tipo, p, label, *_ in campos if p == ph
                )
                novos[ph] = (label, valor)

        if not novos:
            return

        linhas = "\n".join(f"  • {label}: {valor!r}" for _, (label, valor) in novos.items())
        if messagebox.askyesno(
            "Salvar sugestões",
            f"Deseja salvar para sugestões futuras?\n\n{linhas}",
        ):
            for ph, (_, valor) in novos.items():
                historico.setdefault(ph, []).insert(0, valor)
            _salvar_historico(historico)
            for ph, combo in self._combos.items():
                combo.configure(values=historico.get(ph, []))

    def _on_erro(self, msg):
        self._set_loading(False)
        self._set_status(f"Erro: {msg}", "error")
        messagebox.showerror("Erro ao exportar", msg)

    # ------------------------------------------------------------------ Helpers

    def _set_loading(self, ativo):
        if ativo:
            self._progress.pack(side="left", padx=8)
            self._progress.start()
        else:
            self._progress.stop()
            self._progress.pack_forget()

    def _set_status(self, msg, tipo="info"):
        cores = {"success": "#2fa843", "error": "#e74c3c", "warning": "#e67e22", "info": "gray55"}
        self._status.configure(text=msg, text_color=cores.get(tipo, "gray55"))


# ---------------------------------------------------------------------------
# Sub-página: Outros Cargos
# ---------------------------------------------------------------------------
class _SubPaginaOutrosCargos(ctk.CTkFrame):

    _NIVEIS_CURSO = [
        "Ensino Fundamental",
        "Ensino Médio",
        "Curso Técnico",
        "Graduação",
        "Pós-Graduação Lato Sensu",
        "Mestrado",
        "Doutorado",
    ]

    # (porcentagem, porcentagem_extenso)
    _NIVEL_PORCENTAGEM = {
        "Ensino Fundamental":       ("5%",  "(cinco porcento)"),
        "Ensino Médio":             ("10%", "(dez porcento)"),
        "Curso Técnico":            ("10%", "(dez porcento)"),
        "Graduação":                ("15%", "(quinze porcento)"),
        "Pós-Graduação Lato Sensu": ("15%", "(quinze porcento)"),
        "Mestrado":                 ("15%", "(quinze porcento)"),
        "Doutorado":                ("15%", "(quinze porcento)"),
    }

    _ENTRY_PLACEHOLDERS = {"{{NUMERO_DOC}}", "{{NOME}}", "{{CURSO}}", "{{INSTITUICAO}}"}
    _DATE_PLACEHOLDERS  = {"{{DATA_INICIO}}"}

    _LABELS = {
        "{{NUMERO_DOC}}":       "Número do Documento",
        "{{NOME}}":             "Nome",
        "{{DATA_INICIO}}":      "Data de Início no Cargo",
        "{{NUMERO_MATRICULA}}": "Número de Matrícula",
        "{{CURSO}}":            "Nome do Curso",
        "{{INSTITUICAO}}":      "Instituição",
    }

    def __init__(self, parent):
        super().__init__(parent, corner_radius=0, fg_color="transparent")
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)

        self._getters:      dict[str, callable]        = {}
        self._date_widgets: dict[str, _DateEntry]      = {}
        self._matricula_w:  _MatriculaEntry | None     = None
        self._combos:       dict[str, ctk.CTkComboBox] = {}
        self._auto_labels:  dict[str, ctk.CTkLabel]    = {}

        self._build_toolbar()
        self._build_form()
        self._build_statusbar()

    # ------------------------------------------------------------------ UI

    def _build_toolbar(self):
        bar = ctk.CTkFrame(self, height=55, fg_color=("gray82", "gray18"))
        bar.grid(row=0, column=0, sticky="ew")
        bar.grid_propagate(False)

        self._progress = ctk.CTkProgressBar(bar, mode="indeterminate", width=120, height=8)

        ctk.CTkButton(
            bar, text="💾 Exportar .docx", width=150, command=self._exportar
        ).pack(side="right", padx=15, pady=10)

    def _build_form(self):
        scroll = ctk.CTkScrollableFrame(self, fg_color="transparent")
        scroll.grid(row=1, column=0, sticky="nsew", padx=24, pady=(16, 0))
        scroll.grid_columnconfigure(1, weight=1)

        historico = _carregar_historico(_HISTORICO_OUTRO_CARGOS_PATH)
        cargos = list(_carregar_cargos().keys()) or ["— Nenhum cargo cadastrado —"]
        self._cargos_data = _carregar_cargos()
        _sem_cargo = "— Nenhum cargo cadastrado —"

        grupos = [
            ("Documento", [
                ("entry",          "{{NUMERO_DOC}}",          "Número do Documento",    None),
            ]),
            ("Dados do Servidor", [
                ("entry",          "{{NOME}}",                "Nome",                   None),
                ("cargo_dd",       "{{CARGO}}",               "Cargo",                  cargos),
                ("date",           "{{DATA_INICIO}}",         "Data de Início no Cargo",None),
                ("matricula",      "{{NUMERO_MATRICULA}}",    "Número de Matrícula",    None),
            ]),
            ("Dados do Cargo", [
                ("auto", "{{PADRAO_CARGO}}",     "Padrão do Cargo",    "—"),
                ("auto", "{{SINTESE_CARGO}}",    "Síntese do Cargo",   "—"),
                ("auto", "{{CARGA_HORARIA}}",    "Carga Horária",      "—"),
                ("auto", "{{REQUISITOS_CARGO}}", "Requisitos do Cargo","—"),
            ]),
            ("Habilitação", [
                ("nivel_curso_dd", "{{NIVEL_CURSO}}",          "Nível do Curso",         None),
                ("combobox",       "{{CURSO}}",                "Nome do Curso",          historico.get("{{CURSO}}", [])),
                ("combobox",       "{{INSTITUICAO}}",          "Instituição",            historico.get("{{INSTITUICAO}}", [])),
                ("auto",           "{{PORCENTAGEM}}",          "Porcentagem",            "5%"),
                ("auto",           "{{PORCENTAGEM_EXTENSO}}", "Porcentagem por Extenso","(cinco porcento)"),
            ]),
            ("Data do Parecer", [
                ("dropdown", "{{DATA_DIA}}", "Dia", (_DIAS,     _DIA_ATUAL)),
                ("dropdown", "{{DATA_MES}}", "Mês", (_MESES_PT, _MES_ATUAL)),
                ("dropdown", "{{DATA_ANO}}", "Ano", (_ANOS,     _ANO_ATUAL)),
            ]),
        ]

        row = 0
        for grupo, campos in grupos:
            ctk.CTkLabel(
                scroll, text=grupo.upper(),
                font=("Arial", 10, "bold"), text_color="gray50",
            ).grid(row=row, column=0, columnspan=2, sticky="w", pady=(16, 4))
            row += 1

            for tipo, ph, label, extra in campos:
                ctk.CTkLabel(
                    scroll, text=label, font=("Arial", 12), anchor="w"
                ).grid(row=row, column=0, sticky="w", padx=(8, 16), pady=4)

                if tipo == "entry":
                    w = ctk.CTkEntry(scroll, height=32, font=("Arial", 12))
                    w.grid(row=row, column=1, sticky="ew", pady=4)
                    self._getters[ph] = w.get

                elif tipo == "date":
                    w = _DateEntry(scroll, height=32, font=("Arial", 12))
                    w.grid(row=row, column=1, sticky="ew", pady=4)
                    self._getters[ph] = w.get
                    self._date_widgets[ph] = w

                elif tipo == "matricula":
                    w = _MatriculaEntry(scroll, height=32, font=("Arial", 12))
                    w.grid(row=row, column=1, sticky="ew", pady=4)
                    self._getters[ph] = w.get
                    self._matricula_w = w

                elif tipo == "combobox":
                    sugestoes = extra if extra else []
                    w = ctk.CTkComboBox(
                        scroll, values=sugestoes, height=32, font=("Arial", 12)
                    )
                    w.set("")
                    w.grid(row=row, column=1, sticky="ew", pady=4)
                    self._getters[ph] = w.get
                    self._combos[ph] = w

                elif tipo == "dropdown":
                    opcoes, default = extra
                    w = ctk.CTkOptionMenu(
                        scroll, values=opcoes, height=32, font=("Arial", 12)
                    )
                    w.set(default)
                    w.grid(row=row, column=1, sticky="ew", pady=4)
                    self._getters[ph] = w.get
                    _bind_scroll_dropdown(w, opcoes)

                elif tipo == "cargo_dd":
                    w = ctk.CTkOptionMenu(
                        scroll, values=extra, height=32, font=("Arial", 12),
                        command=self._on_cargo_alterado,
                    )
                    w.set(extra[0])
                    w.grid(row=row, column=1, sticky="ew", pady=4)
                    self._getters[ph] = lambda _w=w: "" if _w.get() == _sem_cargo else _w.get()
                    _bind_scroll_dropdown(w, extra)

                elif tipo == "nivel_curso_dd":
                    w = ctk.CTkOptionMenu(
                        scroll, values=self._NIVEIS_CURSO, height=32, font=("Arial", 12),
                        command=self._on_nivel_curso_alterado,
                    )
                    w.set(self._NIVEIS_CURSO[0])
                    w.grid(row=row, column=1, sticky="ew", pady=4)
                    self._getters[ph] = w.get
                    _bind_scroll_dropdown(w, self._NIVEIS_CURSO)

                elif tipo == "auto":
                    lbl = ctk.CTkLabel(
                        scroll, text=extra, font=("Arial", 12),
                        text_color="gray55", anchor="w",
                    )
                    lbl.grid(row=row, column=1, sticky="w", padx=6, pady=4)
                    self._auto_labels[ph] = lbl
                    self._getters[ph] = lambda _l=lbl: _l.cget("text")

                row += 1

    def _build_statusbar(self):
        self._status = ctk.CTkLabel(self, text="", font=("Arial", 11), text_color="gray55")
        self._status.grid(row=2, column=0, padx=25, pady=(4, 8), sticky="w")

    # ------------------------------------------------------------------ Callbacks

    def _on_cargo_alterado(self, cargo: str):
        dados = self._cargos_data.get(cargo, {})
        for ph, chave in (
            ("{{PADRAO_CARGO}}",     "padrao"),
            ("{{SINTESE_CARGO}}",    "sintese"),
            ("{{CARGA_HORARIA}}",    "carga_horaria"),
            ("{{REQUISITOS_CARGO}}", "requisitos"),
        ):
            if ph in self._auto_labels:
                self._auto_labels[ph].configure(text=dados.get(chave, "—"))

    def preencher_teste(self):
        _preencher_widgets_teste(self)
        entries = _listar_entries(self)
        if len(entries) > 1:
            nome, _, _ = _dados_pessoais_teste()
            _preencher_entry(entries[0], str(random.randint(1000, 9999)))
            _preencher_entry(entries[1], nome)

    def _on_nivel_curso_alterado(self, nivel: str):
        pct, pct_ext = self._NIVEL_PORCENTAGEM.get(nivel, ("", ""))
        if "{{PORCENTAGEM}}" in self._auto_labels:
            self._auto_labels["{{PORCENTAGEM}}"].configure(text=pct)
        if "{{PORCENTAGEM_EXTENSO}}" in self._auto_labels:
            self._auto_labels["{{PORCENTAGEM_EXTENSO}}"].configure(text=pct_ext)

    # ------------------------------------------------------------------ Export

    def _exportar(self):
        if not os.path.exists(_TEMPLATE_OUTRO_CARGOS):
            messagebox.showerror("Erro", f"Template não encontrado:\n{_TEMPLATE_OUTRO_CARGOS}")
            return

        substituicoes = {ph: getter() for ph, getter in self._getters.items()}
        substituicoes.update(getattr(self, "_obter_assinatura", lambda: {})())

        vazios         = [ph for ph in self._ENTRY_PLACEHOLDERS if not substituicoes.get(ph, "").strip()]
        datas_invalidas = [ph for ph, w in self._date_widgets.items() if not w.is_valid()]
        matricula_invalida = self._matricula_w is not None and not self._matricula_w.is_valid()

        if vazios or datas_invalidas or matricula_invalida:
            msgs = []
            for ph, label in self._LABELS.items():
                if ph in vazios or ph in datas_invalidas:
                    sufixo = " (data inválida)" if ph in datas_invalidas else ""
                    msgs.append(label + sufixo)
            if matricula_invalida and "{{NUMERO_MATRICULA}}" not in vazios:
                msgs.append("Número de Matrícula (somente dígitos, máx 4)")
            messagebox.showwarning(
                "Campos inválidos",
                "Corrija os campos antes de exportar:\n• " + "\n• ".join(msgs),
            )
            return

        destino = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Documento Word", "*.docx")],
            initialfile="Relatorio Outro Cargos_editado.docx",
        )
        if not destino:
            return

        self._set_loading(True)
        threading.Thread(
            target=self._thread_exportar,
            args=(_TEMPLATE_OUTRO_CARGOS, substituicoes, destino),
            daemon=True,
        ).start()

    def _thread_exportar(self, origem, subs, destino):
        try:
            doc = Document(origem)
            for para in doc.paragraphs:
                _substituir_paragrafo(para, subs)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            _substituir_paragrafo(para, subs)
            doc.save(destino)
            self.after(0, lambda: self._on_exportado(destino))
        except Exception as e:
            msg = str(e)
            self.after(0, lambda: self._on_erro(msg))

    def _on_exportado(self, destino):
        self._set_loading(False)
        self._set_status(f"Exportado com sucesso: {os.path.basename(destino)}", "success")
        self.after(200, self._verificar_novas_sugestoes)

    def _verificar_novas_sugestoes(self):
        historico = _carregar_historico(_HISTORICO_OUTRO_CARGOS_PATH)
        novos = {}
        for ph, combo in self._combos.items():
            valor = combo.get().strip()
            if valor and valor not in historico.get(ph, []):
                novos[ph] = (self._LABELS.get(ph, ph), valor)

        if not novos:
            return

        linhas = "\n".join(f"  • {label}: {valor!r}" for _, (label, valor) in novos.items())
        if messagebox.askyesno(
            "Salvar sugestões",
            f"Deseja salvar para sugestões futuras?\n\n{linhas}",
        ):
            for ph, (_, valor) in novos.items():
                historico.setdefault(ph, []).insert(0, valor)
            _salvar_historico(historico, _HISTORICO_OUTRO_CARGOS_PATH)
            for ph, combo in self._combos.items():
                combo.configure(values=historico.get(ph, []))

    def _on_erro(self, msg):
        self._set_loading(False)
        self._set_status(f"Erro: {msg}", "error")
        messagebox.showerror("Erro ao exportar", msg)

    # ------------------------------------------------------------------ Helpers

    def _set_loading(self, ativo):
        if ativo:
            self._progress.pack(side="left", padx=8)
            self._progress.start()
        else:
            self._progress.stop()
            self._progress.pack_forget()

    def _set_status(self, msg, tipo="info"):
        cores = {"success": "#2fa843", "error": "#e74c3c", "warning": "#e67e22", "info": "gray55"}
        self._status.configure(text=msg, text_color=cores.get(tipo, "gray55"))


class _SubPaginaCTS(ctk.CTkFrame):

    def __init__(self, parent):
        super().__init__(parent, corner_radius=0, fg_color="transparent")
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)
        self._getters = {}
        self._periodos = []
        self._cargo_data = _carregar_cargos()
        self._build_toolbar()
        self._build_form()
        self._build_statusbar()

    def _build_toolbar(self):
        bar = ctk.CTkFrame(self, height=55, fg_color=("gray82", "gray18"))
        bar.grid(row=0, column=0, sticky="ew")
        bar.grid_propagate(False)
        self._progress = ctk.CTkProgressBar(bar, mode="indeterminate", width=120, height=8)
        ctk.CTkButton(bar, text="💾 Exportar .docx", width=150, command=self._exportar).pack(
            side="right", padx=15, pady=10
        )

    def _criar_data(self, parent, row, titulo):
        widgets = []
        ctk.CTkLabel(parent, text=titulo, font=("Arial", 12), anchor="w").grid(
            row=row, column=0, sticky="w", padx=(8, 16), pady=4
        )
        frame = ctk.CTkFrame(parent, fg_color="transparent")
        frame.grid(row=row, column=1, sticky="ew", pady=4)
        larguras = [58, 118, 72]
        valores = [(_DIAS, _DIA_ATUAL), (_MESES_PT, _MES_ATUAL), (_ANOS, _ANO_ATUAL)]
        for col, (opcoes, padrao) in enumerate(valores):
            frame.grid_columnconfigure(col, weight=0)
            w = ctk.CTkOptionMenu(frame, values=opcoes, width=larguras[col], height=32)
            w.set(padrao)
            w.grid(row=0, column=col, padx=(0 if col == 0 else 4, 0))
            _bind_scroll_dropdown(w, opcoes)
            widgets.append(w)
        return widgets

    def _build_form(self):
        scroll = ctk.CTkScrollableFrame(self, fg_color="transparent")
        scroll.grid(row=1, column=0, sticky="nsew", padx=24, pady=(16, 0))
        scroll.grid_columnconfigure(1, weight=1)
        row = 0

        campos = [("entry", "{{NUMERO}}", "Número"), ("letters", "{{NOME}}", "Nome"),
                  ("genero", "{{GENERO_1}}", "Gênero"), ("digits", "{{RG}}", "RG"),
                  ("cpf", "{{CPF}}", "CPF")]
        ctk.CTkLabel(scroll, text="DADOS DO SERVIDOR", font=("Arial", 10, "bold"), text_color="gray50").grid(
            row=row, column=0, columnspan=2, sticky="w", pady=(16, 4)
        )
        row += 1
        for tipo, placeholder, label in campos:
            ctk.CTkLabel(scroll, text=label, font=("Arial", 12), anchor="w").grid(
                row=row, column=0, sticky="w", padx=(8, 16), pady=4
            )
            if tipo == "entry":
                w = ctk.CTkEntry(scroll, height=32, font=("Arial", 12))
                self._getters[placeholder] = w.get
            elif tipo == "letters":
                w = _LettersEntry(scroll, height=32, font=("Arial", 12))
                self._getters[placeholder] = lambda _w=w: _w.get().upper()
            elif tipo == "digits":
                w = _DigitsEntry(scroll, height=32, font=("Arial", 12))
                self._getters[placeholder] = w.get
            elif tipo == "cpf":
                w = _CpfEntry(scroll, height=32, font=("Arial", 12))
                self._getters[placeholder] = w.get
            elif tipo == "genero":
                w = ctk.CTkOptionMenu(scroll, values=["Masculino", "Feminino"], height=32)
                w.set("Masculino")
                self._getters["{{GENERO_1}}"] = lambda _w=w: "a" if _w.get() == "Feminino" else ""
                self._getters["{{GENERO_AO}}"] = lambda _w=w: "a" if _w.get() == "Feminino" else "o"
                _bind_scroll_dropdown(w, ["Masculino", "Feminino"])
            w.grid(row=row, column=1, sticky="ew", pady=4)
            row += 1

        ctk.CTkLabel(scroll, text="PERÍODO", font=("Arial", 10, "bold"), text_color="gray50").grid(
            row=row, column=0, columnspan=2, sticky="w", pady=(16, 4)
        )
        row += 1
        self._periodos_frame = ctk.CTkFrame(scroll, fg_color="transparent")
        self._periodos_frame.grid(row=row, column=0, columnspan=2, sticky="ew")
        self._periodos_frame.grid_columnconfigure(1, weight=1)
        self._adicionar_periodo()
        row += 1
        ctk.CTkLabel(scroll, text="Data do Documento", font=("Arial", 12), anchor="w").grid(
            row=row, column=0, sticky="w", padx=(8, 16), pady=4
        )
        frame = ctk.CTkFrame(scroll, fg_color="transparent")
        frame.grid(row=row, column=1, sticky="ew", pady=4)
        for col, (opcoes, padrao) in enumerate([(_DIAS, _DIA_ATUAL), (_MESES_PT, _MES_ATUAL), (_ANOS, _ANO_ATUAL)]):
            frame.grid_columnconfigure(col, weight=0)
            w = ctk.CTkOptionMenu(frame, values=opcoes, width=[58, 118, 72][col], height=32)
            w.set(padrao)
            w.grid(row=0, column=col, padx=(0 if col == 0 else 4, 0))
            self._getters[f"{{{{DATA_DOC_{col}}}}}"] = w.get
            _bind_scroll_dropdown(w, opcoes)

    def _adicionar_periodo(self):
        periodo = {"frame": None, "inicio": [], "fim": [], "sem_fim": None,
                   "matricula": None, "portaria": None, "cargo": None}
        frame = ctk.CTkFrame(self._periodos_frame, border_width=1, border_color=("gray70", "gray35"))
        frame.grid(row=len(self._periodos), column=0, columnspan=2, sticky="ew", pady=(0, 10))
        frame.grid_columnconfigure(1, weight=1)
        periodo["frame"] = frame
        numero = len(self._periodos) + 1
        ctk.CTkLabel(frame, text=f"Período {numero}", font=("Arial", 12, "bold")).grid(
            row=0, column=0, sticky="w", padx=8, pady=(8, 4)
        )
        if numero > 1:
            ctk.CTkButton(frame, text="-", width=34, command=lambda p=periodo: self._remover_periodo(p)).grid(
                row=0, column=1, sticky="e", padx=8, pady=(6, 2)
            )
        periodo["inicio"] = self._criar_data(frame, 1, "Data Início")
        periodo["fim"] = self._criar_data(frame, 2, "Data Fim")
        periodo["sem_fim"] = ctk.CTkCheckBox(
            frame,
            text="Sem data fim",
            command=lambda p=periodo: self._alternar_data_fim(p),
        )
        ctk.CTkButton(frame, text="+", width=34, command=lambda p=periodo: self._adicionar_a_partir_de(p)).grid(
            row=3, column=1, sticky="e", padx=8, pady=4
        )
        ctk.CTkLabel(frame, text="Matrícula", font=("Arial", 12), anchor="w").grid(
            row=4, column=0, sticky="w", padx=(8, 16), pady=4
        )
        periodo["matricula"] = _DigitsEntry(frame, height=32, font=("Arial", 12))
        periodo["matricula"].grid(row=4, column=1, sticky="ew", pady=4)
        ctk.CTkLabel(frame, text="Portaria", font=("Arial", 12), anchor="w").grid(
            row=5, column=0, sticky="w", padx=(8, 16), pady=4
        )
        periodo["portaria"] = ctk.CTkEntry(frame, height=32, font=("Arial", 12))
        periodo["portaria"].grid(row=5, column=1, sticky="ew", pady=4)
        ctk.CTkLabel(frame, text="Cargo", font=("Arial", 12), anchor="w").grid(
            row=6, column=0, sticky="w", padx=(8, 16), pady=4
        )
        opcoes = list(self._cargo_data) or ["— Nenhum cargo cadastrado —"]
        periodo["cargo"] = ctk.CTkOptionMenu(frame, values=opcoes, height=32)
        periodo["cargo"].set(opcoes[0])
        periodo["cargo"].grid(row=6, column=1, sticky="ew", pady=4)
        _bind_scroll_dropdown(periodo["cargo"], opcoes)
        self._periodos.append(periodo)
        self._atualizar_controles_periodos()

    def _adicionar_a_partir_de(self, periodo):
        if periodo["sem_fim"].get():
            messagebox.showwarning("Período inválido", "Desmarque 'Sem data fim' antes de adicionar outro período.")
            return
        self._adicionar_periodo()

    def _remover_periodo(self, periodo):
        if len(self._periodos) == 1:
            return
        periodo["frame"].destroy()
        self._periodos.remove(periodo)
        for row, item in enumerate(self._periodos):
            item["frame"].grid_configure(row=row)
        self._atualizar_controles_periodos()

    def _alternar_data_fim(self, periodo):
        estado = ctk.DISABLED if periodo["sem_fim"].get() else ctk.NORMAL
        for widget in periodo["fim"]:
            widget.configure(state=estado)

    def _atualizar_controles_periodos(self):
        ultimo = self._periodos[-1]
        for periodo in self._periodos:
            if periodo is ultimo:
                periodo["sem_fim"].grid(row=3, column=1, sticky="w", padx=6, pady=4)
            else:
                periodo["sem_fim"].grid_forget()
                periodo["sem_fim"].deselect()
            self._alternar_data_fim(periodo)

    def _obter_data_widgets(self, widgets):
        mes = _MESES_PT.index(widgets[1].get()) + 1
        return datetime(int(widgets[2].get()), mes, int(widgets[0].get()))

    def preencher_teste(self):
        _preencher_widgets_teste(self)
        entries = _listar_entries(self)
        nome, rg, cpf = _dados_pessoais_teste()
        for indice, valor in enumerate((str(random.randint(100000000, 999999999)), nome, rg, cpf)):
            if indice < len(entries):
                _preencher_entry(entries[indice], valor)
        while len(self._periodos) < 2:
            self._adicionar_periodo()
        inicio = datetime(random.randint(2022, 2025), random.randint(1, 12), random.randint(1, 28))
        for indice, periodo in enumerate(self._periodos):
            data_inicio = inicio + timedelta(days=indice * 100)
            data_fim = data_inicio + timedelta(days=60)
            _definir_data_teste(periodo["inicio"], data_inicio)
            _definir_data_teste(periodo["fim"], data_fim)
            periodo["sem_fim"].deselect()
            _preencher_entry(periodo["matricula"], str(1000 + indice))

    def _coletar_periodos(self):
        resultado = []
        for indice, periodo in enumerate(self._periodos):
            numero = indice + 1
            try:
                inicio = self._obter_data_widgets(periodo["inicio"])
            except (ValueError, IndexError):
                raise ValueError(f"Período {numero}: Data Início inválida") from None
            sem_fim = bool(periodo["sem_fim"].get())
            try:
                fim = datetime.now() if sem_fim else self._obter_data_widgets(periodo["fim"])
            except (ValueError, IndexError):
                raise ValueError(f"Período {numero}: Data Fim inválida") from None
            if fim < inicio:
                raise ValueError(f"Período {numero}: Data Fim anterior à Data Início")
            if sem_fim and indice != len(self._periodos) - 1:
                raise ValueError("Somente o último período pode estar sem data fim")
            cargo = periodo["cargo"].get()
            if cargo.startswith("—"):
                cargo = ""
            resultado.append({"periodo": f"{_data_por_extenso(inicio)} até {'a presente data da emissão desta declaração' if sem_fim else _data_por_extenso(fim)}",
                              "dias": (fim - inicio).days, "matricula": periodo["matricula"].get(),
                              "portaria": periodo["portaria"].get(), "cargo": cargo,
                              "sintese": self._cargo_data.get(cargo, {}).get("sintese", "")})
        return resultado

    def _build_statusbar(self):
        self._status = ctk.CTkLabel(self, text="", font=("Arial", 11), text_color="gray55")
        self._status.grid(row=2, column=0, padx=25, pady=(4, 8), sticky="w")

    def _exportar(self):
        if not os.path.exists(_TEMPLATE_CTS):
            messagebox.showerror("Erro", f"Template não encontrado:\n{_TEMPLATE_CTS}")
            return
        subs = {ph: getter() for ph, getter in self._getters.items()}
        subs.update(getattr(self, "_obter_assinatura", lambda: {})())
        problemas = []
        labels = {
            "{{NUMERO}}": "Número",
            "{{NOME}}": "Nome",
            "{{RG}}": "RG",
        }
        problemas.extend(
            label for placeholder, label in labels.items()
            if not subs.get(placeholder, "").strip()
        )
        cpf = self._getters["{{CPF}}"]()
        cpf_digitos = cpf.replace(".", "").replace("-", "")
        if not cpf_digitos or not cpf_digitos.isdigit() or len(cpf_digitos) != 11:
            problemas.append("CPF (deve conter 11 dígitos)")
        for indice, periodo in enumerate(self._periodos):
            numero = indice + 1
            if not periodo["matricula"].get().strip():
                problemas.append(f"Período {numero}: Matrícula")
            if not periodo["portaria"].get().strip():
                problemas.append(f"Período {numero}: Portaria")
            cargo = periodo["cargo"].get()
            if cargo.startswith("—"):
                problemas.append(f"Período {numero}: Cargo")
        try:
            periodos = self._coletar_periodos()
        except ValueError as erro:
            problemas.append(str(erro))
            periodos = []
        try:
            data_doc = self._obter_data("DATA_DOC")
        except (ValueError, IndexError):
            problemas.append("Data do Documento inválida")
            data_doc = None
        if problemas:
            messagebox.showwarning(
                "Campos inválidos",
                "Corrija os campos antes de exportar:\n• " + "\n• ".join(problemas),
            )
            return
        for indice, periodo in enumerate(periodos):
            if not periodo["sintese"].strip():
                problemas.append(f"Período {indice + 1}: Síntese do cargo")
        if problemas:
            messagebox.showwarning(
                "Campos inválidos",
                "Corrija os campos antes de exportar:\n• " + "\n• ".join(problemas),
            )
            return
        total_dias = sum(periodo["dias"] for periodo in periodos)
        detalhes_periodos = "\r".join(
            f"Durante o período de {periodo['periodo']} sob a matrícula nº {periodo['matricula']} "
            f"no cargo de {periodo['cargo']}, conforme Portaria nº {periodo['portaria']}, "
            f"exercendo as atribuições de: {periodo['sintese']} "
            f"Totalizando {periodo['dias']} ({_numero_por_extenso(periodo['dias'])}) dias de serviços prestados a este Município."
            for periodo in periodos
        )
        texto_periodos = detalhes_periodos
        if len(periodos) > 1:
            texto_periodos += (
                f"\rTotalizando assim {total_dias} "
                f"({_numero_por_extenso(total_dias)}) dias de serviços prestados "
                "a este Município durante os períodos mencionados."
            )
        subs.update({
            "{{BLOCO_PERIODOS}}": texto_periodos,
            "{{DATA}}": _data_por_extenso(data_doc),
            "{{NEGRITO_CTS}}": [
                *(periodo["periodo"] for periodo in periodos),
                *(periodo["cargo"] for periodo in periodos),
                *([f"{total_dias} ({_numero_por_extenso(total_dias)}) dias de serviços prestados"]
                  if len(periodos) > 1 else []),
            ],
        })
        destino = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Documento Word", "*.docx")], initialfile="Relatorio cts_editado.docx")
        if not destino:
            return
        self._set_loading(True)
        threading.Thread(target=self._thread_exportar, args=(subs, destino), daemon=True).start()

    def _obter_data(self, prefix):
        valores = [self._getters[f"{{{{{prefix}_{col}}}}}"]() for col in range(3)]
        mes = _MESES_PT.index(valores[1]) + 1
        return datetime(int(valores[2]), mes, int(valores[0]))

    def _thread_exportar(self, subs, destino):
        word = doc = None
        com_inicializado = False
        try:
            import pythoncom
            import win32com.client
            pythoncom.CoInitialize()
            com_inicializado = True
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            doc = word.Documents.Open(
                os.path.abspath(_TEMPLATE_CTS),
                ConfirmConversions=False,
                ReadOnly=False,
                AddToRecentFiles=False,
            )
            def substituir_no_intervalo(intervalo, texto, valor):
                inicio = intervalo.Start
                while inicio < doc.Content.End:
                    fim = doc.Content.End
                    alvo = doc.Range(inicio, fim)
                    localizar = alvo.Find
                    localizar.ClearFormatting()
                    localizar.Text = texto
                    localizar.Forward = True
                    localizar.Wrap = 0
                    if not localizar.Execute():
                        break
                    alvo.Text = str(valor)
                    novo_inicio = alvo.End
                    if novo_inicio <= inicio:
                        novo_inicio = inicio + 1
                    inicio = novo_inicio

            def substituir(texto, valor):
                substituir_no_intervalo(doc.Content.Duplicate, texto, valor)
                for secao in doc.Sections:
                    for cabecalho in secao.Headers:
                        substituir_no_intervalo(cabecalho.Range.Duplicate, texto, valor)
                    for rodape in secao.Footers:
                        substituir_no_intervalo(rodape.Range.Duplicate, texto, valor)

            def aplicar_negrito(texto):
                if not texto:
                    return
                intervalo = doc.Content.Duplicate
                localizar = intervalo.Find
                localizar.ClearFormatting()
                localizar.Text = texto
                localizar.Forward = True
                localizar.Wrap = 0
                while localizar.Execute():
                    intervalo.Font.Bold = True
                    intervalo.SetRange(intervalo.End, doc.Content.End)

            texto_periodos = subs.pop("{{BLOCO_PERIODOS}}", None)
            textos_negrito = subs.pop("{{NEGRITO_CTS}}", [])
            if texto_periodos is not None:
                marcador_inicio = doc.Content.Duplicate
                marcador_inicio.Find.ClearFormatting()
                marcador_inicio.Find.Text = "{{PERIODO}}"
                marcador_inicio.Find.Forward = True
                marcador_inicio.Find.Wrap = 0
                marcador_fim = doc.Content.Duplicate
                marcador_fim.Find.ClearFormatting()
                marcador_fim.Find.Text = "{{PERIODO_EXTENSO}}"
                marcador_fim.Find.Forward = True
                marcador_fim.Find.Wrap = 0
                if not marcador_inicio.Find.Execute() or not marcador_fim.Find.Execute():
                    raise ValueError("O fim do trecho dos períodos não foi encontrado no template CTS.")

                inicio = marcador_inicio.Start - len("durante o período de ")
                if inicio < 0 or marcador_fim.End < inicio:
                    raise ValueError("O trecho dos períodos não foi encontrado no template CTS.")
                doc.Range(inicio, marcador_fim.End).Text = texto_periodos
                sufixo = "dias de serviços prestados a este Município."
                texto_documento = doc.Content.Text
                sufixo_inicio = texto_documento.rfind(sufixo)
                if sufixo_inicio >= 0:
                    doc.Range(sufixo_inicio, sufixo_inicio + len(sufixo)).Text = ""

            for chave, valor in subs.items():
                substituir(chave, valor)
            restantes = [chave for chave in (
                "{{NUMERO}}", "{{NOME}}", "{{GENERO_1}}", "{{RG}}",
                "{{GENERO_AO}}", "{{CPF}}", "{{MATRICULA}}", "{{CARGO}}",
                "{{PORTARIA}}", "{{SINTESE}}", "{{PERIODO}}", "{{PERIODO_EXTENSO}}",
                "{{DATA}}", "{{NOME_ASSINATURA}}", "{{CARGO ASSINATURA}}",
                "{{PORTARIA ASSINATURA}}",
            ) if chave in doc.Content.Text]
            if restantes:
                raise ValueError("Placeholders não substituídos no CTS: " + ", ".join(restantes))
            for texto in textos_negrito:
                aplicar_negrito(texto)
            nome = subs["{{NOME}}"]
            if nome:
                aplicar_negrito(nome)
            doc.SaveAs2(os.path.abspath(destino), FileFormat=16)
            self.after(0, lambda: self._on_exportado(destino))
        except Exception as e:
            msg = str(e)
            self.after(0, lambda: self._on_erro(msg))
        finally:
            if doc is not None:
                try:
                    doc.Close(False)
                except Exception:
                    pass
            if word is not None:
                try:
                    quit_word = getattr(word, "Quit", None)
                    if callable(quit_word):
                        quit_word()
                except Exception:
                    pass
            if com_inicializado:
                pythoncom.CoUninitialize()

    def _on_exportado(self, destino):
        self._set_loading(False)
        self._status.configure(text=f"Exportado com sucesso: {os.path.basename(destino)}", text_color="#2fa843")

    def _on_erro(self, msg):
        self._set_loading(False)
        self._status.configure(text=f"Erro: {msg}", text_color="#e74c3c")
        messagebox.showerror("Erro ao exportar CTS", msg)

    def _set_loading(self, ativo):
        if ativo:
            self._progress.pack(side="left", padx=8)
            self._progress.start()
        else:
            self._progress.stop()
            self._progress.pack_forget()


class _SubPaginaEstagiario(_SubPaginaCTS):

    def __init__(self, parent):
        self._atribuicoes = _carregar_atribuicoes()
        super().__init__(parent)

    def preencher_teste(self):
        _preencher_widgets_teste(self)
        entries = _listar_entries(self)
        nome, rg, cpf = _dados_pessoais_teste()
        for indice, valor in enumerate((nome, cpf, rg)):
            if indice < len(entries):
                _preencher_entry(entries[indice], valor)
        while len(self._periodos) < 2:
            self._adicionar_periodo()
        inicio = datetime(random.randint(2022, 2025), random.randint(1, 12), random.randint(1, 28))
        for indice, periodo in enumerate(self._periodos):
            data_inicio = inicio + timedelta(days=indice * 100)
            data_fim = data_inicio + timedelta(days=60)
            _definir_data_teste(periodo["inicio"], data_inicio)
            _definir_data_teste(periodo["fim"], data_fim)
            periodo["sem_fim"].deselect()
            _preencher_entry(periodo["local"], _local_teste(indice))
            _preencher_entry(periodo["matricula"], str(random.randint(1000, 9999)))

    def _build_form(self):
        scroll = ctk.CTkScrollableFrame(self, fg_color="transparent")
        scroll.grid(row=1, column=0, sticky="nsew", padx=24, pady=(16, 0))
        scroll.grid_columnconfigure(1, weight=1)
        row = 0

        ctk.CTkLabel(scroll, text="DADOS DO ESTAGIÁRIO", font=("Arial", 10, "bold"), text_color="gray50").grid(
            row=row, column=0, columnspan=2, sticky="w", pady=(16, 4)
        )
        row += 1
        campos = (
            ("{{NOME}}", "Nome", "letters"),
            ("{{CPF}}", "CPF", "cpf"),
            ("{{RG}}", "RG", "digits"),
        )
        for placeholder, label, tipo in campos:
            ctk.CTkLabel(scroll, text=label, font=("Arial", 12), anchor="w").grid(
                row=row, column=0, sticky="w", padx=(8, 16), pady=4
            )
            if tipo == "letters":
                widget = _LettersEntry(scroll, height=32, font=("Arial", 12))
            elif tipo == "cpf":
                widget = _CpfEntry(scroll, height=32, font=("Arial", 12))
            elif tipo == "digits":
                widget = _DigitsEntry(scroll, height=32, font=("Arial", 12))
            else:
                widget = ctk.CTkEntry(scroll, height=32, font=("Arial", 12))
            widget.grid(row=row, column=1, sticky="ew", pady=4)
            self._getters[placeholder] = lambda _w=widget: _w.get().strip()
            row += 1

        ctk.CTkLabel(scroll, text="PERÍODOS", font=("Arial", 10, "bold"), text_color="gray50").grid(
            row=row, column=0, columnspan=2, sticky="w", pady=(16, 4)
        )
        row += 1
        self._periodos_frame = ctk.CTkFrame(scroll, fg_color="transparent")
        self._periodos_frame.grid(row=row, column=0, columnspan=2, sticky="ew")
        self._periodos_frame.grid_columnconfigure(1, weight=1)
        self._adicionar_periodo()
        row += 1
        ctk.CTkLabel(scroll, text="Data do Documento", font=("Arial", 12), anchor="w").grid(
            row=row, column=0, sticky="w", padx=(8, 16), pady=4
        )
        data_frame = ctk.CTkFrame(scroll, fg_color="transparent")
        data_frame.grid(row=row, column=1, sticky="ew", pady=4)
        for coluna, (opcoes, padrao) in enumerate(((_DIAS, _DIA_ATUAL), (_MESES_PT, _MES_ATUAL), (_ANOS, _ANO_ATUAL))):
            widget = ctk.CTkOptionMenu(data_frame, values=opcoes, width=[58, 118, 72][coluna], height=32)
            widget.set(padrao)
            widget.grid(row=0, column=coluna, padx=(0 if coluna == 0 else 4, 0))
            self._getters[f"{{{{DATA_DOC_{coluna}}}}}"] = widget.get
            _bind_scroll_dropdown(widget, opcoes)

    def _adicionar_periodo(self):
        periodo = {"frame": None, "inicio": [], "fim": [], "sem_fim": None,
               "local": None, "matricula": None, "atividades": None}
        frame = ctk.CTkFrame(self._periodos_frame, border_width=1, border_color=("gray70", "gray35"))
        frame.grid(row=len(self._periodos), column=0, columnspan=2, sticky="ew", pady=(0, 10))
        frame.grid_columnconfigure(1, weight=1)
        periodo["frame"] = frame
        numero = len(self._periodos) + 1
        ctk.CTkLabel(frame, text=f"Período {numero}", font=("Arial", 12, "bold")).grid(
            row=0, column=0, sticky="w", padx=8, pady=(8, 4)
        )
        if numero > 1:
            ctk.CTkButton(frame, text="-", width=34, command=lambda p=periodo: self._remover_periodo(p)).grid(
                row=0, column=1, sticky="e", padx=8, pady=(6, 2)
            )
        periodo["inicio"] = self._criar_data(frame, 1, "Data Início")
        periodo["fim"] = self._criar_data(frame, 2, "Data Fim")
        periodo["sem_fim"] = ctk.CTkCheckBox(frame, text="Sem data fim", command=lambda p=periodo: self._alternar_data_fim(p))
        ctk.CTkButton(frame, text="+", width=34, command=lambda p=periodo: self._adicionar_a_partir_de(p)).grid(
            row=3, column=1, sticky="e", padx=8, pady=4
        )
        ctk.CTkLabel(frame, text="Local", font=("Arial", 12), anchor="w").grid(
            row=4, column=0, sticky="w", padx=(8, 16), pady=4
        )
        periodo["local"] = ctk.CTkEntry(frame, height=32, font=("Arial", 12))
        periodo["local"].grid(row=4, column=1, sticky="ew", pady=4)
        ctk.CTkLabel(frame, text="Matrícula", font=("Arial", 12), anchor="w").grid(
            row=5, column=0, sticky="w", padx=(8, 16), pady=4
        )
        periodo["matricula"] = _DigitsEntry(frame, height=32, font=("Arial", 12))
        periodo["matricula"].grid(row=5, column=1, sticky="ew", pady=4)
        ctk.CTkLabel(frame, text="Atividades", font=("Arial", 12), anchor="w").grid(
            row=6, column=0, sticky="w", padx=(8, 16), pady=4
        )
        opcoes = [item["nome"] for item in self._atribuicoes] or ["— Nenhuma atribuição cadastrada —"]
        periodo["atividades"] = ctk.CTkOptionMenu(frame, values=opcoes, height=32)
        periodo["atividades"].set(opcoes[0])
        periodo["atividades"].grid(row=6, column=1, sticky="ew", pady=4)
        _bind_scroll_dropdown(periodo["atividades"], opcoes)
        self._periodos.append(periodo)
        self._atualizar_controles_periodos()

    def _coletar_periodos(self):
        resultado = []
        descricoes = {item["nome"]: item["descricao"] for item in self._atribuicoes}
        for indice, periodo in enumerate(self._periodos):
            numero = indice + 1
            try:
                inicio = self._obter_data_widgets(periodo["inicio"])
            except (ValueError, IndexError):
                raise ValueError(f"Período {numero}: Data Início inválida") from None
            sem_fim = bool(periodo["sem_fim"].get())
            try:
                fim = datetime.now() if sem_fim else self._obter_data_widgets(periodo["fim"])
            except (ValueError, IndexError):
                raise ValueError(f"Período {numero}: Data Fim inválida") from None
            if fim < inicio:
                raise ValueError(f"Período {numero}: Data Fim anterior à Data Início")
            if sem_fim and indice != len(self._periodos) - 1:
                raise ValueError("Somente o último período pode estar sem data fim")
            atividade = periodo["atividades"].get()
            if atividade.startswith("—"):
                atividade = ""
            resultado.append({
                "periodo": f"{_data_por_extenso(inicio)} até {'a presente data da emissão desta declaração' if sem_fim else _data_por_extenso(fim)}",
                "dias": (fim - inicio).days,
                "local": periodo["local"].get().strip(),
                "matricula": periodo["matricula"].get(),
                "atividade": atividade,
                "descricao": descricoes.get(atividade, ""),
            })
        return resultado

    def _exportar(self):
        if not os.path.exists(_TEMPLATE_ESTAGIO):
            messagebox.showerror("Erro", f"Template não encontrado:\n{_TEMPLATE_ESTAGIO}")
            return
        subs = {ph: getter() for ph, getter in self._getters.items()}
        subs.update(getattr(self, "_obter_assinatura", lambda: {})())
        problemas = [label for ph, label in (("{{NOME}}", "Nome"),) if not subs[ph]]
        cpf = subs["{{CPF}}"].replace(".", "").replace("-", "")
        if len(cpf) != 11 or not cpf.isdigit():
            problemas.append("CPF (deve conter 11 dígitos)")
        if not subs["{{RG}}"].strip():
            problemas.append("RG")
        for indice, periodo in enumerate(self._periodos):
            if not periodo["local"].get().strip():
                problemas.append(f"Período {indice + 1}: Local")
            if not periodo["matricula"].get().strip():
                problemas.append(f"Período {indice + 1}: Matrícula")
            if periodo["atividades"].get().startswith("—"):
                problemas.append(f"Período {indice + 1}: Atividades")
        try:
            periodos = self._coletar_periodos()
        except ValueError as erro:
            problemas.append(str(erro))
            periodos = []
        try:
            data_doc = self._obter_data_estagio()
        except (ValueError, IndexError):
            problemas.append("Data do Documento inválida")
            data_doc = None
        if problemas:
            messagebox.showwarning("Campos inválidos", "Corrija os campos antes de exportar:\n• " + "\n• ".join(problemas))
            return
        total_dias = sum(item["dias"] for item in periodos)
        bloco = "\r".join(
            f"{item['local']} sob matrícula n° {item['matricula']}, desenvolvendo as atividades a seguir: {item['descricao']}, durante o período de {item['periodo']}."
            for item in periodos
        )
        subs["{{TOTALIZADOR}}"] = f"{total_dias} ({_numero_por_extenso(total_dias)})"
        subs["{{DATA}}"] = _data_por_extenso(data_doc)
        subs["{{NEGRITO_ESTAGIO}}"] = [
            subs["{{NOME}}"],
            subs["{{TOTALIZADOR}}"],
            *(item["periodo"] for item in periodos),
        ]
        destino = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Documento Word", "*.docx")], initialfile="Relatorio estagio_editado.docx")
        if not destino:
            return
        self._set_loading(True)
        threading.Thread(target=self._thread_exportar_estagio, args=(subs, bloco, destino), daemon=True).start()

    def _obter_data_estagio(self):
        valores = [self._getters[f"{{{{DATA_DOC_{coluna}}}}}"]() for coluna in range(3)]
        return datetime(int(valores[2]), _MESES_PT.index(valores[1]) + 1, int(valores[0]))

    def _thread_exportar_estagio(self, subs, bloco, destino):
        word = doc = None
        com_inicializado = False
        try:
            import pythoncom
            import win32com.client
            pythoncom.CoInitialize()
            com_inicializado = True
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            doc = word.Documents.Open(os.path.abspath(_TEMPLATE_ESTAGIO), ConfirmConversions=False, ReadOnly=False, AddToRecentFiles=False, NoEncodingDialog=True)
            inicio = doc.Content.Duplicate
            inicio.Find.Text = "{{LOCAL}}"
            if not inicio.Find.Execute():
                raise ValueError("O marcador {{LOCAL}} não foi encontrado no template de estágio.")
            fim = doc.Content.Duplicate
            fim.Find.Text = "{{PERIODO}}"
            if not fim.Find.Execute():
                raise ValueError("O trecho dos períodos não foi encontrado no template de estágio.")
            bloco_range = doc.Range(inicio.Start, fim.End)
            bloco_range.Text = bloco
            bloco_range.Font.Bold = False
            textos_negrito = subs.pop("{{NEGRITO_ESTAGIO}}", [])
            for chave, valor in subs.items():
                intervalo = doc.Content.Duplicate
                localizar = intervalo.Find
                localizar.Text = chave
                localizar.Forward = True
                localizar.Wrap = 0
                while localizar.Execute():
                    intervalo.Text = str(valor)
                    intervalo.Collapse(0)
            for texto in textos_negrito:
                if not texto:
                    continue
                intervalo = doc.Content.Duplicate
                localizar = intervalo.Find
                localizar.Text = texto
                localizar.Forward = True
                localizar.Wrap = 0
                while localizar.Execute():
                    intervalo.Font.Bold = True
                    intervalo.Collapse(0)
            doc.SaveAs2(os.path.abspath(destino), FileFormat=16)
            self.after(0, lambda: self._on_exportado(destino))
        except Exception as erro:
            mensagem = str(erro)
            self.after(0, lambda: self._on_erro(mensagem))
        finally:
            if doc is not None:
                doc.Close(False)
            if word is not None:
                word.Quit()
            if com_inicializado:
                pythoncom.CoUninitialize()

    def _set_loading(self, ativo):
        if ativo:
            self._progress.pack(side="left", padx=8)
            self._progress.start()
        else:
            self._progress.stop()
            self._progress.pack_forget()


# ---------------------------------------------------------------------------
# Utilitário: substituição preservando formatação dos runs
# ---------------------------------------------------------------------------
def _substituir_paragrafo(para, subs: dict):
    texto = "".join(r.text for r in para.runs)
    novo = texto
    for chave, valor in subs.items():
        novo = novo.replace(chave, valor)
    if novo != texto and para.runs:
        run0 = para.runs[0]
        rPr = run0._r.find(qn('w:rPr'))
        for r in para.runs:
            r.text = ""
        partes = novo.split('\n')
        run0.text = partes[0]
        prev_r = run0._r
        for parte in partes[1:]:
            br_r = OxmlElement('w:r')
            if rPr is not None:
                br_r.append(copy.deepcopy(rPr))
            br_r.append(OxmlElement('w:br'))
            prev_r.addnext(br_r)
            text_r = OxmlElement('w:r')
            if rPr is not None:
                text_r.append(copy.deepcopy(rPr))
            t = OxmlElement('w:t')
            t.text = parte
            if parte.startswith(' ') or parte.endswith(' '):
                t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            text_r.append(t)
            br_r.addnext(text_r)
            prev_r = text_r
